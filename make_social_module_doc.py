#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Genera Modulo_Comunicaciones_MTTSIA.docx
Presentación del Módulo de Comunicaciones y Redes Sociales
para la videollamada del Ministerio de Transportes — Guinea Ecuatorial
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_font(run, name='Calibri', size=11, bold=False, italic=False, color=None):
    run.font.name   = name
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, level=1, color=(0,70,127)):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(*color)
    return p

def para(text='', bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return p

def shade_row(row, fill_hex='003366'):
    for cell in row.cells:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  fill_hex)
        tcPr.append(shd)

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0]
    shade_row(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            set_font(run, size=10, bold=True, color=(255,255,255))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row_data in enumerate(rows):
        row = t.add_row()
        if ri % 2 == 1:
            shade_row(row, 'EBF3FB')
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                set_font(run, size=10)
    if col_widths:
        for row in t.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Cm(w)
    doc.add_paragraph()
    return t

def step_box(number, title, steps):
    """Numbered step block with shaded title."""
    p_title = doc.add_paragraph()
    r_num  = p_title.add_run(f'  Paso {number}  ')
    set_font(r_num, size=11, bold=True, color=(255,255,255))
    r_num.font.highlight_color = None
    # shade the title paragraph manually
    from docx.oxml import OxmlElement as OE
    pPr = p_title._p.get_or_add_pPr()
    shd = OE('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  '1F497D')
    pPr.append(shd)
    r_title = p_title.add_run(f'  {title}')
    set_font(r_title, size=11, bold=True, color=(255,255,255))
    for s in steps:
        bp = doc.add_paragraph(style='List Bullet')
        set_font(bp.add_run(s), size=10)

def responsible_card(role, name, responsibilities):
    """Highlighted responsible-person card."""
    p = doc.add_paragraph()
    r1 = p.add_run(f'{role}:  ')
    set_font(r1, size=11, bold=True, color=(0,70,127))
    r2 = p.add_run(name)
    set_font(r2, size=11, bold=True, color=(192,0,0))
    for resp in responsibilities:
        bp = doc.add_paragraph(style='List Bullet')
        set_font(bp.add_run(resp), size=10)


# ═══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run('MINISTERIO DE TRANSPORTES, TELECOMUNICACIONES')
set_font(r, size=13, bold=True, color=(0,70,127))

title_p2 = doc.add_paragraph()
title_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = title_p2.add_run('CORREOS, INFRAESTRUCTURA E INTERNET DE GUINEA ECUATORIAL')
set_font(r2, size=13, bold=True, color=(0,70,127))

doc.add_paragraph()
divider = doc.add_paragraph('─' * 65)
divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

main_title = doc.add_paragraph()
main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = main_title.add_run('MÓDULO DE COMUNICACIONES Y REDES SOCIALES')
set_font(r3, size=20, bold=True, color=(0,70,127))

doc.add_paragraph()
sub_title = doc.add_paragraph()
sub_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = sub_title.add_run('Guía de Presentación para Videollamada')
set_font(r4, size=13, bold=True, color=(31,73,125))

doc.add_paragraph()
sub_title2 = doc.add_paragraph()
sub_title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r5 = sub_title2.add_run('Centro de Mando Ministerial — Plataforma Digital con IA')
set_font(r5, size=11, italic=True, color=(127,127,127))

doc.add_paragraph()
doc.add_paragraph()

meta_rows = [
    ('Documento:', 'Presentación del Módulo de Comunicaciones'),
    ('Versión:',   '1.0 — Edición para Videollamada'),
    ('Fecha:',     datetime.date.today().strftime('%d de %B de %Y')),
    ('Preparado para:', 'Videollamada con el Ministerio — Revisión de Funcionalidades'),
    ('Audiencia:',  'Ministro, Técnico Ministerial, Responsables de Área'),
    ('Estado del módulo:', '✅ COMPLETADO Y DESPLEGADO EN PRODUCCIÓN'),
]
for label, value in meta_rows:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_l = p.add_run(f'{label}  ')
    set_font(r_l, size=11, bold=True, color=(0,70,127))
    r_v = p.add_run(value)
    set_font(r_v, size=11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# ÍNDICE RÁPIDO
# ═══════════════════════════════════════════════════════════════════════════════
heading('ÍNDICE DE CONTENIDOS', 1)
index_items = [
    ('1', 'Visión General del Módulo de Comunicaciones',               'Pág. 3'),
    ('2', 'Responsables de Área',                                      'Pág. 3'),
    ('3', 'Sección XI — Generación de Contenido de Prensa (IA)',       'Pág. 4'),
    ('4', 'Sección XII — Módulo Multimedia con IA',                    'Pág. 6'),
    ('5', 'Sección XIV — Comunicación por WhatsApp',                   'Pág. 7'),
    ('6', 'Sección XIII — Asistente IA Ministerial',                   'Pág. 8'),
    ('7', 'Guion de Demostración para la Videollamada',                'Pág. 9'),
    ('8', 'Preguntas Frecuentes (FAQ)',                                 'Pág. 10'),
]
add_table(['#', 'Sección', 'Referencia'], index_items, col_widths=[0.8, 13.0, 2.0])

# ═══════════════════════════════════════════════════════════════════════════════
# 1. VISIÓN GENERAL
# ═══════════════════════════════════════════════════════════════════════════════
heading('1. VISIÓN GENERAL DEL MÓDULO DE COMUNICACIONES', 1)
para(
    'El Módulo de Comunicaciones es el conjunto de herramientas digitales del Centro de Mando '
    'Ministerial destinadas a gestionar, generar y distribuir contenido oficial del Ministerio '
    'de Transportes, Telecomunicaciones, Correos, Infraestructura e Internet de Guinea Ecuatorial.',
    size=11
)
doc.add_paragraph()
para(
    'Este módulo integra Inteligencia Artificial (IA) para automatizar y agilizar las '
    'comunicaciones institucionales, reduciendo el tiempo de redacción y garantizando '
    'la coherencia y calidad del mensaje ministerial en todos los canales.',
    size=11
)
doc.add_paragraph()

# What it covers
overview_data = [
    ('XI',   'Prensa y Contenido',  'Artículos, borradores y posts para redes sociales generados con IA'),
    ('XII',  'Multimedia',          'Transcripción, traducción y dictamen técnico de audio y video'),
    ('XIII', 'Asistente IA',        'Chat inteligente para redacción y consultas institucionales'),
    ('XIV',  'Bot WhatsApp',         'Chatbot que responde consultas del personal por WhatsApp mediante comandos'),
]
add_table(['Sección', 'Sub-Módulo', 'Función Principal'], overview_data, col_widths=[1.8, 3.5, 11.0])

para(
    'Todos los sub-módulos están completamente desarrollados, desplegados en el servidor de '
    'producción (http://157.230.178.118) y listos para uso inmediato.',
    size=11, bold=True, color=(0,128,0)
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 2. RESPONSABLES DE ÁREA
# ═══════════════════════════════════════════════════════════════════════════════
heading('2. RESPONSABLES DE ÁREA — MÓDULO DE COMUNICACIONES', 1)
para(
    'Cada sub-módulo debe contar con un responsable designado por el Ministerio. '
    'A continuación se presenta la estructura de responsabilidades recomendada:',
    size=11
)
doc.add_paragraph()

resp_data = [
    ('Director de Comunicaciones',
     'ADMIN / GABINETE',
     'Responsable del módulo completo',
     [
        'Aprobar y publicar artículos y contenido de prensa',
        'Supervisar la generación de contenido con IA',
        'Autorizar posts para redes sociales',
        'Gestionar la agenda de publicaciones',
     ]),
    ('Jefe de Prensa',
     'GABINETE / REVISOR',
     'Responsable de Sección XI (Contenido de Prensa)',
     [
        'Generar propuestas de artículos semanales con IA',
        'Revisar y editar borradores generados',
        'Crear y copiar posts para redes sociales',
        'Coordinar publicaciones por sector ministerial',
     ]),
    ('Técnico de Multimedia',
     'GABINETE / REVISOR',
     'Responsable de Sección XII (Multimedia)',
     [
        'Subir y gestionar archivos de audio y video',
        'Supervisar la transcripción automática con IA',
        'Solicitar traducciones de transcripciones',
        'Revisar dictámenes técnicos generados',
     ]),
    ('Responsable de Comunicación Digital',
     'GABINETE',
     'Responsable de Secciones XIII–XIV (Asistente IA + Bot WhatsApp)',
     [
        'Gestionar el Asistente IA para redacción oficial',
        'Activar y supervisar la conexión del Bot WhatsApp',
        'Gestionar el teléfono del Ministerio vinculado al bot',
        'Verificar el registro de actividad del bot',
        'Registrar los números WhatsApp del personal en el sistema',
     ]),
]

for role, permission, title, duties in resp_data:
    p_header = doc.add_paragraph()
    pPr = p_header._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  '003366')
    pPr.append(shd)
    r_role = p_header.add_run(f'  {role}')
    set_font(r_role, size=12, bold=True, color=(255,255,255))
    r_sep = p_header.add_run('   |   ')
    set_font(r_sep, size=11, color=(200,220,255))
    r_perm = p_header.add_run(f'Rol requerido: {permission}')
    set_font(r_perm, size=10, italic=True, color=(200,220,255))

    p_title = doc.add_paragraph()
    r_t = p_title.add_run(f'  {title}')
    set_font(r_t, size=11, bold=True, color=(0,70,127))

    for duty in duties:
        bp = doc.add_paragraph(style='List Bullet')
        set_font(bp.add_run(duty), size=10)
    doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 3. SECCIÓN XI — PRENSA Y CONTENIDO
# ═══════════════════════════════════════════════════════════════════════════════
heading('3. SECCIÓN XI — GENERACIÓN DE CONTENIDO DE PRENSA CON IA', 1)

para(
    'Este sub-módulo permite al equipo de prensa generar artículos institucionales completos '
    'y publicaciones para redes sociales de forma automática utilizando Inteligencia Artificial '
    '(GPT-4o de OpenAI), en español y adaptados al contexto del Ministerio.',
    size=11
)
doc.add_paragraph()

heading('3.1 Funcionalidades Disponibles', 2)
features_XI = [
    ('Propuestas Semanales con IA',
     'El sistema genera automáticamente 6 propuestas de artículos (una por cada sector ministerial) con título, introducción y esquema de contenido.'),
    ('Borradores Completos',
     'Desde una propuesta, la IA redacta un artículo completo de más de 600 palabras con estructura formal: introducción, desarrollo por secciones y conclusión.'),
    ('Posts para Redes Sociales',
     'El sistema genera automáticamente dos versiones del post: versión corta para Twitter/X (máximo 280 caracteres) y versión extendida para Facebook/Instagram con hashtags y emojis.'),
    ('Gestión de Artículos',
     'Crear, editar, revisar, publicar y archivar artículos. Vista paginada con 10 artículos por página. Filtros por estado (Borrador, Pendiente, Publicado).'),
    ('Control de Roles',
     'Solo ADMIN y GABINETE pueden generar contenido con IA. REVISOR puede editar borradores. LECTOR solo puede ver artículos publicados.'),
]
feat_data = [(f, d) for f, d in features_XI]
add_table(['Funcionalidad', 'Descripción'], feat_data, col_widths=[5.0, 10.5])

heading('3.2 Sectores Cubiertos por la IA', 2)
sectors = [
    ('Transportes',           'Infraestructura vial, carreteras, transporte público'),
    ('Puertos',               'Puertos marítimos, logística naval, comercio exterior'),
    ('Telecomunicaciones',    'Redes de telecomunicaciones, conectividad, regulación'),
    ('Inteligencia Artificial','Innovación tecnológica, digitalización gubernamental'),
    ('Correos',               'Servicio postal nacional, distribución y logística'),
    ('Regulación',            'Marco normativo, legislación sectorial, fiscalización'),
]
add_table(['Sector', 'Áreas de Contenido'], sectors, col_widths=[4.5, 11.0])

heading('3.3 Cómo Usar — Guía Paso a Paso', 2)

steps_XI_generate = [
    'Ir al menú lateral y seleccionar "Contenido" (ícono de periódico)',
    'En la parte superior derecha, hacer clic en el botón azul "Generar con IA"',
    'El sistema tarda 15–30 segundos en generar las 6 propuestas (una por sector)',
    'Aparece un panel con las 6 propuestas: título, sector, introducción y esquema',
    'Desplazarse (scroll) para ver todas las propuestas en el panel',
]
step_box(1, 'Generar Propuestas Semanales', steps_XI_generate)

steps_XI_draft = [
    'En el panel de propuestas, localizar la propuesta deseada',
    'Hacer clic en el botón "Redactar" de esa propuesta',
    'La IA genera un borrador completo (+600 palabras) en 15–20 segundos',
    'El borrador se guarda automáticamente como artículo en estado "Borrador"',
    'El artículo aparece en la lista de la página "Contenido"',
    'Hacer clic en el artículo para abrirlo, editarlo o enviarlo a revisión',
]
step_box(2, 'Generar Borrador de Artículo Completo', steps_XI_draft)

steps_XI_social = [
    'En la lista de artículos, localizar el artículo deseado',
    'Hacer clic en el botón "Red Social" (ícono de compartir) del artículo',
    'La IA genera en segundos: versión Twitter y versión Facebook',
    'Usar el botón "Copiar" para copiar el texto al portapapeles',
    'Pegar directamente en la red social correspondiente',
]
step_box(3, 'Generar Post para Redes Sociales', steps_XI_social)

steps_XI_publish = [
    'Abrir el artículo en estado "Borrador" o "En Revisión"',
    'Revisar y editar el contenido si es necesario',
    'Hacer clic en "Publicar" (solo ADMIN y GABINETE)',
    'El artículo queda visible para todos los usuarios de la plataforma',
]
step_box(4, 'Publicar un Artículo', steps_XI_publish)
doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 4. SECCIÓN XII — MÓDULO MULTIMEDIA
# ═══════════════════════════════════════════════════════════════════════════════
heading('4. SECCIÓN XII — MÓDULO MULTIMEDIA CON IA', 1)
para(
    'El módulo multimedia permite al Ministerio gestionar archivos de audio y video, '
    'obtener transcripciones automáticas, traducciones a múltiples idiomas y dictámenes '
    'técnicos oficiales generados con Inteligencia Artificial.',
    size=11
)
doc.add_paragraph()

heading('4.1 Funcionalidades Disponibles', 2)
features_XII = [
    ('Carga de Archivos',           'Sube archivos de audio (MP3, WAV) y video (MP4, WebM) hasta 50 MB'),
    ('Transcripción Automática',    'OpenAI Whisper transcribe el contenido hablado a texto en segundos'),
    ('Traducción con IA',           'Traduce la transcripción a Francés, Inglés, Ruso, Chino o Portugués'),
    ('Dictamen Técnico Ministerial', 'La IA genera un informe oficial con 5 secciones: Resumen Ejecutivo, Relevancia Institucional, Análisis de Contenido, Riesgos y Oportunidades, Recomendaciones'),
    ('Reproductor Integrado',       'Reproduce el archivo directamente en la plataforma sin descargar'),
    ('Búsqueda y Filtros',          'Buscar por nombre, fecha, idioma o estado de transcripción'),
]
feat_data2 = [(f, d) for f, d in features_XII]
add_table(['Funcionalidad', 'Descripción'], feat_data2, col_widths=[5.0, 10.5])

heading('4.2 Cómo Usar — Guía Paso a Paso', 2)

steps_XII_upload = [
    'Ir al menú lateral y seleccionar "Multimedia"',
    'Hacer clic en "Subir Archivo" (botón superior derecho)',
    'Seleccionar el archivo de audio o video desde el equipo',
    'Añadir título, descripción y el idioma del archivo',
    'Hacer clic en "Subir" y esperar la confirmación',
]
step_box(1, 'Subir un Archivo de Audio o Video', steps_XII_upload)

steps_XII_transcribe = [
    'Abrir el archivo multimedia haciendo clic en su nombre',
    'En el panel de detalle, localizar la sección "Transcripción"',
    'Si no hay transcripción, hacer clic en "Transcribir con IA"',
    'El sistema procesa el audio (10–30 segundos según duración)',
    'El texto transcrito aparece en el panel de detalle',
    'La transcripción queda guardada permanentemente en el sistema',
]
step_box(2, 'Transcribir Audio/Video a Texto', steps_XII_transcribe)

steps_XII_translate = [
    'Con la transcripción ya disponible, localizar el selector de idioma',
    'Seleccionar el idioma destino: Francés, Inglés, Ruso, Chino o Portugués',
    'Hacer clic en "Traducir"',
    'La traducción aparece debajo de la transcripción original',
    'Ambas versiones (original + traducción) quedan guardadas',
]
step_box(3, 'Traducir la Transcripción', steps_XII_translate)

steps_XII_opinion = [
    'Con la transcripción disponible, localizar el botón "Generar Dictamen Técnico"',
    'La IA analiza el contenido y genera un informe estructurado en 5 secciones',
    'El dictamen aparece en el panel (fondo naranja para identificarlo fácilmente)',
    'El informe puede copiarse o imprimirse para uso oficial',
]
step_box(4, 'Generar Dictamen Técnico Ministerial', steps_XII_opinion)
doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 5. SECCIÓN XIV — WHATSAPP BOT
# ═══════════════════════════════════════════════════════════════════════════════
heading('5. SECCIÓN XIV — BOT WHATSAPP MINISTERIAL', 1)
para(
    'El módulo "Bot WhatsApp" permite al Ministerio activar un chatbot oficial que '
    'responde automáticamente a consultas del personal a través de WhatsApp. '
    'El personal envía comandos de texto al número del bot y recibe respuestas '
    'instantáneas sobre documentos, plazos y agenda, sin necesidad de abrir el sistema.',
    size=11
)
doc.add_paragraph()

# Info box clarifying what this module is
p_info = doc.add_paragraph()
pPr_info = p_info._p.get_or_add_pPr()
shd_info = OxmlElement('w:shd')
shd_info.set(qn('w:val'),   'clear')
shd_info.set(qn('w:color'), 'auto')
shd_info.set(qn('w:fill'),  'E8F5E9')
pPr_info.append(shd_info)
r_info = p_info.add_run(
    '  ACCESO AL MÓDULO: Menú lateral → "Bot WhatsApp" (ícono de mensaje verde).  '
    'El módulo permite conectar un número de teléfono real al sistema mediante código QR, '
    'igual que vincular WhatsApp Web en un ordenador.'
)
set_font(r_info, size=10, color=(27,94,32))
doc.add_paragraph()

features_XIV = [
    ('Conexión por Código QR',
     'Vincula un número de WhatsApp al sistema escaneando un código QR, igual que WhatsApp Web'),
    ('Estado de Conexión en Tiempo Real',
     'Panel que muestra si el bot está Conectado, Iniciando, Esperando QR o Desconectado'),
    ('Número del Bot Visible',
     'Una vez conectado, muestra el número de teléfono oficial del Ministerio vinculado al bot'),
    ('Comandos de Consulta Rápida',
     'El personal envía comandos por WhatsApp y el bot responde automáticamente (ver tabla de comandos)'),
    ('Registro de Actividad',
     'Historial de todos los mensajes recibidos y respuestas enviadas por el bot'),
    ('Usuarios Registrados',
     'Lista de usuarios del sistema con su número de WhatsApp vinculado'),
    ('Reconexión Automática',
     'Botón para reconectar el bot si se desconecta sin necesidad de soporte técnico'),
]
add_table(['Funcionalidad', 'Descripción'], [(f, d) for f, d in features_XIV], col_widths=[5.0, 10.5])

heading('5.1 Comandos Disponibles del Bot', 2)
para(
    'El personal del Ministerio envía estos comandos como mensajes de WhatsApp al número del bot:',
    size=11
)
commands_data = [
    ('ayuda',          'Ver la lista completa de comandos disponibles'),
    ('pendientes',     'Ver todos los documentos en estado Pendiente'),
    ('urgentes',       'Ver documentos marcados como urgentes'),
    ('plazos',         'Ver los plazos que vencen hoy'),
    ('agenda',         'Ver los eventos de la agenda de hoy'),
    ('estado [num]',   'Consultar el estado de un documento por su número (ej: estado 025-MT-038)'),
    ('buscar [texto]', 'Buscar documentos por palabra clave (ej: buscar convenio transporte)'),
    ('estadisticas',   'Ver un resumen general del sistema (totales por estado)'),
]
add_table(['Comando', 'Acción'], commands_data, col_widths=[4.5, 11.0])

heading('5.2 Cómo Activar el Bot — Guía Paso a Paso', 2)

steps_XIV_connect = [
    'Ir al menú lateral y hacer clic en "Bot WhatsApp"',
    'Si el estado es "Desconectado", hacer clic en el botón "Reconectar"',
    'Esperar entre 5 y 15 segundos hasta que aparezca el código QR en pantalla',
    'Abrir WhatsApp en el teléfono oficial del Ministerio',
    'Ir a: Ajustes (tres puntos) → Dispositivos vinculados → Vincular dispositivo',
    'Apuntar la cámara al código QR de la pantalla',
    'El estado cambia a "Conectado" y se muestra el número de teléfono vinculado',
]
step_box(1, 'Conectar el Bot por Primera Vez', steps_XIV_connect)

steps_XIV_use = [
    'Desde cualquier teléfono del personal, abrir WhatsApp',
    'Enviar un mensaje al número oficial del bot del Ministerio',
    'Escribir uno de los comandos disponibles (ej: "pendientes")',
    'El bot responde automáticamente en segundos con la información solicitada',
    'El mensaje y la respuesta quedan registrados en "Actividad reciente" del panel',
]
step_box(2, 'Usar el Bot desde un Teléfono', steps_XIV_use)

steps_XIV_users = [
    'Ir a Configuración → Usuarios en el menú lateral',
    'Editar el perfil del usuario deseado',
    'En el campo "WhatsApp", introducir su número con prefijo internacional (ej: +240 222 658 556)',
    'Guardar — ese usuario ahora puede recibir notificaciones automáticas por WhatsApp',
]
step_box(3, 'Registrar el Número WhatsApp de un Usuario', steps_XIV_users)
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 6. SECCIÓN XIII — ASISTENTE IA
# ═══════════════════════════════════════════════════════════════════════════════
heading('6. SECCIÓN XIII — ASISTENTE IA MINISTERIAL', 1)
para(
    'El Asistente IA es un chat inteligente integrado en la plataforma que permite '
    'al equipo ministerial redactar, consultar y analizar documentos oficiales '
    'mediante lenguaje natural, sin necesidad de herramientas externas.',
    size=11
)
doc.add_paragraph()

features_XIII = [
    ('Redacción de Documentos Oficiales', 'Generar decretos, circulares, notas verbales, oficios e informes'),
    ('Consultas Institucionales',         'Hacer preguntas sobre procedimientos, legislación y normativa'),
    ('Análisis de Texto',                 'Resumir, mejorar o corregir textos redactados por el usuario'),
    ('Respuestas en Español Formal',      'Todas las respuestas adaptadas al registro formal ministerial'),
    ('Copiado Rápido',                    'Botón para copiar cualquier respuesta directamente al portapapeles'),
    ('Historial de Conversación',         'El historial se mantiene durante la sesión activa'),
]
add_table(['Funcionalidad', 'Descripción'], [(f, d) for f, d in features_XIII], col_widths=[5.5, 10.0])

heading('6.1 Cómo Usar — Ejemplos de Consultas', 2)

examples = [
    ('"Redacta un decreto sobre la regulación de telecomunicaciones en Guinea Ecuatorial"',
     'El asistente genera un decreto completo con estructura oficial'),
    ('"Resume el siguiente texto en 3 puntos clave: [pegar texto]"',
     'Obtiene un resumen conciso del contenido'),
    ('"¿Cuál es el procedimiento para registrar una empresa de telecomunicaciones?"',
     'El asistente responde con el procedimiento en formato de lista'),
    ('"Redacta una nota verbal de agradecimiento al embajador de España"',
     'Genera un texto formal diplomático listo para usar'),
    ('"Traduce este párrafo al francés: [texto en español]"',
     'Traducción instantánea al idioma solicitado'),
]
add_table(['Ejemplo de Consulta', 'Resultado Esperado'], examples, col_widths=[8.0, 7.5])

steps_XIII = [
    'Ir al menú lateral y seleccionar "Asistente IA" (ícono de estrella/chat)',
    'Escribir la consulta o instrucción en el campo de texto inferior',
    'Presionar Enter o hacer clic en el botón de enviar',
    'Leer la respuesta generada en el panel de chat',
    'Usar el botón "Copiar" para copiar el texto generado',
    'Pegar directamente en el documento, correo o sistema que necesite',
]
step_box(1, 'Usar el Asistente IA', steps_XIII)
doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 7. GUION PARA VIDEOLLAMADA
# ═══════════════════════════════════════════════════════════════════════════════
heading('7. GUION DE DEMOSTRACIÓN PARA LA VIDEOLLAMADA', 1)
para(
    'A continuación se presenta el orden de demostración recomendado para la videollamada '
    'de revisión del Módulo de Comunicaciones. Duración estimada: 45–60 minutos.',
    size=11
)
doc.add_paragraph()

demo_data = [
    ('00:00–05:00', '1', 'Presentación del módulo',
     'Mostrar la pantalla principal. Explicar los 4 sub-módulos y sus responsables de área.'),
    ('05:00–15:00', '2', 'Demo Sección XI — Prensa y IA',
     '1. Clic en "Contenido"\n2. Clic en "Generar con IA" → mostrar 6 propuestas\n3. Clic en "Redactar" → ver borrador generado\n4. Clic en "Red Social" → ver post Twitter y Facebook\n5. Demostrar función "Copiar"'),
    ('15:00–25:00', '3', 'Demo Sección XII — Multimedia',
     '1. Clic en "Multimedia"\n2. Subir un archivo de audio/video de prueba\n3. Hacer clic en "Transcribir" → ver resultado\n4. Seleccionar idioma → clic en "Traducir"\n5. Clic en "Generar Dictamen Técnico" → ver informe'),
    ('25:00–32:00', '4', 'Demo Sección XIV — Bot WhatsApp',
     '1. Clic en "Bot WhatsApp" en el menú lateral\n2. Mostrar panel de estado del bot (Conectado / QR)\n3. Si está desconectado: clic "Reconectar" → escanear QR en directo\n4. Mostrar tabla de comandos disponibles\n5. Enviar comando "pendientes" desde un teléfono real y ver respuesta en vivo\n6. Mostrar registro de actividad reciente del bot'),
    ('32:00–40:00', '5', 'Demo Sección XIII — Asistente IA',
     '1. Clic en "Asistente IA"\n2. Escribir: "Redacta una circular sobre horarios de atención"\n3. Ver la respuesta generada\n4. Copiar y pegar en editor de Word\n5. Hacer 1-2 consultas adicionales del técnico ministerial'),
    ('40:00–50:00', '6', 'Preguntas y prueba libre',
     'El técnico ministerial y el equipo del Ministerio prueban el módulo libremente con sus propios casos de uso reales. El desarrollador responde dudas.'),
    ('50:00–60:00', '7', 'Revisión de responsables y cierre',
     'Confirmar los nombres de los responsables de área designados. Revisar accesos y permisos. Planificar el siguiente paso (UAT, hosting, pago).'),
]

for time, num, title, desc in demo_data:
    p_hdr = doc.add_paragraph()
    pPr = p_hdr._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1F497D')
    pPr.append(shd)
    r_time = p_hdr.add_run(f'  {time}  ')
    set_font(r_time, size=10, bold=True, color=(200,220,255))
    r_num = p_hdr.add_run(f'  Bloque {num}:  ')
    set_font(r_num, size=11, bold=True, color=(255,255,255))
    r_title = p_hdr.add_run(title)
    set_font(r_title, size=11, bold=True, color=(255,230,100))

    for line in desc.split('\n'):
        bp = doc.add_paragraph(style='List Bullet')
        set_font(bp.add_run(line.lstrip('0123456789. ')), size=10)
    doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 8. FAQ
# ═══════════════════════════════════════════════════════════════════════════════
heading('8. PREGUNTAS FRECUENTES (FAQ)', 1)

faqs = [
    ('¿Cuánto tarda en generarse un artículo completo?',
     'Entre 15 y 30 segundos. La IA procesa la propuesta y redacta más de 600 palabras con estructura formal completa.'),
    ('¿Se puede editar el artículo generado por la IA?',
     'Sí. Todos los artículos generados se guardan como "Borrador" y pueden editarse libremente antes de publicarse.'),
    ('¿Los posts de redes sociales se publican automáticamente?',
     'No. La plataforma genera el texto del post, pero el responsable debe copiarlo y publicarlo manualmente en la red social correspondiente. Esto es intencional para mantener el control editorial del Ministerio.'),
    ('¿Qué idiomas soporta la transcripción de audio?',
     'El sistema soporta todos los idiomas de OpenAI Whisper: español, francés, inglés, árabe, portugués, chino y más de 50 idiomas adicionales.'),
    ('¿Qué pasa si la transcripción tiene errores?',
     'El texto transcrito es editable. El responsable puede corregir manualmente errores antes de usar la transcripción. La calidad depende de la claridad del audio.'),
    ('¿Desde qué número responde el Bot WhatsApp?',
     'El bot usa el número del teléfono del Ministerio que se vincula escaneando el código QR en la pantalla. Una vez conectado, el propio panel muestra el número activo. El personal envía sus consultas a ese número y el bot responde automáticamente.'),
    ('¿Quién puede usar el Asistente IA?',
     'Todos los usuarios con cuenta activa en la plataforma (ADMIN, GABINETE, REVISOR, LECTOR) pueden usar el Asistente IA. La generación de contenido de prensa (Sección XI) está limitada a ADMIN y GABINETE.'),
    ('¿El contenido generado por la IA es confidencial?',
     'Sí. Todo el contenido se procesa a través de la API de OpenAI con las políticas de privacidad correspondientes. Los datos del Ministerio no se usan para entrenar modelos de IA.'),
    ('¿Qué ocurre si se pierde la conexión durante la generación?',
     'Si la conexión se interrumpe, la operacion puede reintentarse. Los articulos ya guardados no se pierden. Recomendamos una conexión estable de al menos 5 Mbps.'),
    ('¿Se puede limitar qué sectores aparecen en las propuestas?',
     'En la versión actual, los 6 sectores ministeriales estan fijos. Si el Ministerio desea agregar o modificar sectores, puede solicitarlo como mejora en el periodo de soporte.'),
]

for q, a in faqs:
    p_q = doc.add_paragraph()
    r_q = p_q.add_run(f'P: {q}')
    set_font(r_q, size=11, bold=True, color=(0,70,127))
    p_a = doc.add_paragraph()
    r_a = p_a.add_run(f'R: {a}')
    set_font(r_a, size=11)
    doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CIERRE
# ═══════════════════════════════════════════════════════════════════════════════
heading('NOTA FINAL PARA LA VIDEOLLAMADA', 1)
para(
    'Todos los módulos descritos en este documento están completamente operativos '
    'y disponibles en el servidor de producción. El equipo de desarrollo estará '
    'disponible durante la videollamada para responder preguntas técnicas, '
    'realizar ajustes en tiempo real si fuese necesario, y presentar cualquier '
    'funcionalidad adicional del sistema.',
    size=11
)
doc.add_paragraph()

access_data = [
    ('URL del Sistema',     'http://157.230.178.118'),
    ('Estado',              'Activo y operativo (24/7)'),
    ('Soporte disponible',  'Durante y despues de la videollamada'),
    ('Proximos pasos',      'UAT, configuracion de dominio HTTPS, formacion de usuarios'),
]
add_table(['Parámetro', 'Detalle'], access_data, col_widths=[5.0, 10.5])

footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_f = footer_p.add_run(
    'Ministerio de Transportes, Telecomunicaciones, Correos, Infraestructura e Internet | '
    'Guinea Ecuatorial | Centro de Mando Ministerial v1.0 | '
    + datetime.date.today().strftime('%d/%m/%Y')
)
set_font(r_f, size=9, italic=True, color=(127,127,127))

# ═══════════════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════════════
output_path = r'f:\workana_work\AfricaMinistrator\ministerial-command-center\Modulo_Comunicaciones_MTTSIA.docx'
doc.save(output_path)
print(f'Document saved: {output_path}')
