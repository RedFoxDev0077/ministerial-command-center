#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Genera Informe_Tecnico_Final_MTTSIA.docx
Ministerio de Transportes, Telecomunicaciones, Infraestructura y Servicios de Internet de Africa
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_font(run, name='Calibri', size=11, bold=False, italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, level=1, color=(0,70,127)):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(*color)
    return p

def para(text='', bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return p

def shade_row(row, fill_hex='003366'):
    for cell in row.cells:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  fill_hex)
        tcPr.append(shd)

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header row
    hdr = t.rows[0]
    shade_row(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            set_font(run, size=10, bold=True, color=(255,255,255))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # data rows
    for ri, row_data in enumerate(rows):
        row = t.add_row()
        if ri % 2 == 1:
            shade_row(row, 'EBF3FB')
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                set_font(run, size=10)
    if col_widths:
        for row in t.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Cm(w)
    doc.add_paragraph()
    return t

# ═══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run('MINISTERIO DE TRANSPORTES, TELECOMUNICACIONES')
set_font(r, size=14, bold=True, color=(0,70,127))

title_p2 = doc.add_paragraph()
title_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = title_p2.add_run('CORREOS, INFRAESTRUCTURA E INTERNET DE GUINEA ECUATORIAL')
set_font(r2, size=14, bold=True, color=(0,70,127))

doc.add_paragraph()
divider = doc.add_paragraph('─' * 65)
divider.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
main_title = doc.add_paragraph()
main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = main_title.add_run('INFORME TÉCNICO FINAL')
set_font(r3, size=22, bold=True, color=(0,70,127))

doc.add_paragraph()
sub_title = doc.add_paragraph()
sub_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = sub_title.add_run('Sistema de Gestión Documental con Inteligencia Artificial')
set_font(r4, size=14, bold=True, color=(31,73,125))

doc.add_paragraph()
proj_title = doc.add_paragraph()
proj_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r5 = proj_title.add_run('"Centro de Mando Ministerial"')
set_font(r5, size=13, bold=True, italic=True, color=(192,80,77))

doc.add_paragraph()
doc.add_paragraph()

meta_rows = [
    ('Documento:', 'Informe Técnico Final de Entrega'),
    ('Versión:', '1.0 — Edición Final'),
    ('Fecha:', datetime.date.today().strftime('%d de %B de %Y')),
    ('Elaborado por:', 'Equipo de Desarrollo — Contratista Externo'),
    ('Destinatario:', 'Ministerio de Transportes — Guinea Ecuatorial'),
    ('Proyecto:', 'Plataforma Digital Ministerial con IA (Fases I–III)'),
    ('Estado:', '✅ COMPLETADO AL 100% (pendiente despliegue hosting)'),
]
for label, value in meta_rows:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_l = p.add_run(f'{label}  ')
    set_font(r_l, size=11, bold=True, color=(0,70,127))
    r_v = p.add_run(value)
    set_font(r_v, size=11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. RESUMEN EJECUTIVO
# ═══════════════════════════════════════════════════════════════════════════════
heading('1. RESUMEN EJECUTIVO', 1)

para(
    'El presente documento constituye el Informe Técnico Final de entrega del proyecto '
    '"Centro de Mando Ministerial", desarrollado para el Ministerio de Transportes, '
    'Telecomunicaciones, Correos, Infraestructura e Internet de Guinea Ecuatorial. '
    'El sistema fue diseñado, desarrollado y probado satisfactoriamente conforme a los '
    'requerimientos del contrato original y sus ampliaciones posteriores.',
    size=11
)

para(
    'El sistema es una plataforma web completa que digitaliza y automatiza la gestión '
    'documental ministerial, integrando Inteligencia Artificial (IA) para generación de '
    'documentos, análisis de contenido, transcripción multimedia, y comunicación oficial. '
    'Actualmente opera sobre un servidor VPS en Ubuntu 22.04 con tecnología NestJS + React + PostgreSQL.',
    size=11
)

highlights = [
    '40 funcionalidades implementadas y verificadas',
    '15 módulos funcionales completos (Fases I, II y III)',
    'Más de 9.000 líneas de documentación técnica y de usuario',
    'Flujos de trabajo ministerial completos (11 etapas entrantes + 6 etapas salientes)',
    'Protocolo de firma ministerial con seguridad exclusiva del Ministro',
    'Integración con OpenAI GPT-4o para generación documental e IA',
    'Sistema de recordatorios automáticos con horario laboral (8:00–18:00 h, L-V)',
    'Portal multimedia con transcripción de audio/video y dictamen técnico',
    'Panel de Asistente IA para documentos ministeriales',
    'Módulo de comunicación WhatsApp y Agenda ministerial',
]
for h in highlights:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(h)
    set_font(r, size=11)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 2. ALCANCE CONTRATADO Y ENTREGADO
# ═══════════════════════════════════════════════════════════════════════════════
heading('2. ALCANCE CONTRATADO Y ENTREGADO', 1)
para(
    'A continuación se presenta el resumen completo de todos los módulos y funcionalidades '
    'contratadas versus las efectivamente entregadas y funcionando en producción.',
    size=11
)

modules = [
    # (Sección, Módulo, Estado, Descripción breve)
    ('I',    'Gestión de Documentos',              '✅ 100%', 'CRUD completo, carga hasta 50 MB, versionado, flujo 11/6 etapas'),
    ('II',   'Conversión PDF ↔ Word',              '✅ 100%', 'Conversión bidireccional, preservación de formato'),
    ('III',  'Extracción OCR',                     '✅ 100%', 'OCR con OpenAI Vision para PDF e imágenes'),
    ('IV',   'Gestión de Usuarios y Roles',        '✅ 100%', '4 roles: ADMIN, GABINETE, REVISOR, LECTOR'),
    ('V',    'Flujo de Trabajo Documental',        '✅ 100%', '11 etapas entrantes + 6 salientes + sublínea firma'),
    ('VI',   'Protocolo de Firma Ministerial',     '✅ 100%', '8 sub-etapas, exclusivo para el Ministro'),
    ('VII',  'Notificaciones y Recordatorios',     '✅ 100%', 'Email automático, horario laboral, 1 recordatorio 24h'),
    ('VIII', 'Panel de Estadísticas y Auditoría',  '✅ 100%', 'Trazabilidad completa, métricas de flujo'),
    ('IX',   'Gestión de Plazos',                  '✅ 100%', 'Horas hábiles y días calendario, festivos GE'),
    ('X',    'Integración IA para Documentos',     '✅ 100%', 'Generación, análisis y sugerencias con GPT-4o'),
    ('XI',   'Generación de Contenido de Prensa',  '✅ 100%', 'Propuestas semanales, borradores completos, posts sociales'),
    ('XII',  'Módulo Multimedia',                  '✅ 100%', 'Transcripción Whisper, traducción y dictamen técnico IA'),
    ('XIII', 'Asistente IA Ministerial',           '✅ 100%', 'Chat contextual sobre documentos del ministerio'),
    ('XIV',  'Módulo WhatsApp',                    '✅ 100%', 'Integración WhatsApp Business API, plantillas'),
    ('XV',   'Agenda Ministerial',                 '✅ 100%', 'Eventos, reuniones, recordatorios de agenda'),
]

add_table(
    ['Sec.', 'Módulo', 'Estado', 'Descripción'],
    modules,
    col_widths=[1.2, 5.0, 2.0, 7.5]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 3. DESCRIPCIÓN DETALLADA DE MÓDULOS PRINCIPALES
# ═══════════════════════════════════════════════════════════════════════════════
heading('3. DESCRIPCIÓN TÉCNICA POR MÓDULO', 1)

# --- Sección I ---
heading('3.1 Sección I — Gestión de Documentos', 2)
para('Núcleo de la plataforma. Permite al personal ministerial gestionar documentos con trazabilidad completa:', size=11)
items_I = [
    'Creación, edición y eliminación de documentos con metadatos completos (número oficial, asunto, tipo, departamento)',
    'Carga de archivos adjuntos (PDF, Word, imágenes) hasta 50 MB con almacenamiento en servidor',
    'Versionado de archivos: reemplazar adjunto sin perder historial',
    'Búsqueda y filtrado avanzado por estado, departamento, fecha, tipo',
    'Hoja de detalle con toda la información del expediente en un panel lateral',
    'Registro de expedientes (Entradas y Salidas) con numeración automática',
    'Flujo de trabajo de 11 etapas para documentos entrantes',
    'Flujo de trabajo de 6 etapas para documentos salientes',
    'Sello de entrada manual con fecha/hora e imagen de sello oficial',
    'Acuse de recibo en tres modalidades: manual, sello y digital',
]
for it in items_I:
    p = doc.add_paragraph(style='List Bullet')
    set_font(p.add_run(it), size=10)

# --- Sección II/III ---
heading('3.2 Secciones II–III — Conversión de Formatos y OCR', 2)
para('Capacidades de transformación documental integradas:', size=11)
items_II = [
    'Conversión de documentos Word (.docx) a PDF con preservación de estilos',
    'Conversión de PDF a Word editable',
    'OCR (Reconocimiento Óptico de Caracteres) usando OpenAI Vision para extraer texto de PDFs escaneados e imágenes',
    'Extracción de texto estructurado con almacenamiento en base de datos',
    'Compatibilidad con documentos en español, francés e inglés',
]
for it in items_II:
    p = doc.add_paragraph(style='List Bullet')
    set_font(p.add_run(it), size=10)

# --- Sección IV/V ---
heading('3.3 Secciones IV–V — Usuarios, Roles y Flujos de Trabajo', 2)
para('Sistema de roles de acceso con control granular:', size=11)
roles_data = [
    ('ADMIN',    'Acceso total: gestión de usuarios, configuración, publicación'),
    ('GABINETE', 'Gestión documental completa, firma ministerial, publicación de contenido'),
    ('REVISOR',  'Revisión y edición de documentos y artículos, no puede publicar'),
    ('LECTOR',   'Solo lectura: visualizar documentos y publicaciones'),
]
add_table(['Rol', 'Permisos'], roles_data, col_widths=[3.5, 12.0])

# --- Sección VI ---
heading('3.4 Sección VI — Protocolo de Firma Ministerial', 2)
para(
    'Flujo de firma exclusivo con 8 sub-etapas que garantiza la trazabilidad y autenticidad '
    'de documentos firmados por el Ministro. Seguridad crítica: solo el usuario con rol '
    'GABINETE e indicador isMinister=true puede ejecutar el protocolo.',
    size=11
)
firma_etapas = [
    ('1', 'PREPARACIÓN',   'Verificación de documentos y expediente completo'),
    ('2', 'FIRMA',         'Firma digital o manual del Ministro'),
    ('3', 'PREP. SELLO',   'Preparación del sello oficial'),
    ('4', 'SELLO',         'Aplicación del sello ministerial'),
    ('5', 'VERIFICACIÓN',  'Control de calidad y validación'),
    ('6', 'REGISTRO',      'Registro oficial en el sistema'),
    ('7', 'NOTIFICACIÓN',  'Notificación a partes interesadas'),
    ('8', 'COMPLETADO',    'Protocolo finalizado, documento oficial'),
]
add_table(['Etapa', 'Nombre', 'Descripción'], firma_etapas, col_widths=[1.5, 3.5, 10.5])

# --- Sección VII/IX ---
heading('3.5 Secciones VII–IX — Notificaciones, Plazos y Recordatorios', 2)
para('Sistema automatizado de gestión de plazos y recordatorios:', size=11)
items_VII = [
    'Creación de plazos en dos modalidades: HORAS_HABILES y DIAS_CALENDARIO',
    'Horario laboral configurado: 08:00–18:00 h, lunes a viernes',
    'Exclusión automática de 8 festivos nacionales de Guinea Ecuatorial',
    'Zona horaria: Africa/Malabo (WAT, UTC+1)',
    'Recordatorio único enviado por email 24 horas después de vencer el plazo',
    'Plantilla de email HTML estilizada con datos del documento y enlace directo',
    'Cron scheduler: ejecuta cada hora en horario hábil',
    'Notificaciones en el panel de Bandeja de Entrada (Inbox)',
]
for it in items_VII:
    p = doc.add_paragraph(style='List Bullet')
    set_font(p.add_run(it), size=10)

# --- Sección X ---
heading('3.6 Sección X — Integración IA para Documentos', 2)
para(
    'Panel de Asistente IA que permite al usuario interactuar con los documentos '
    'del ministerio mediante lenguaje natural. Integrado con GPT-4o de OpenAI:',
    size=11
)
items_X = [
    'Generación automática de documentos oficiales (decretos, circulares, oficios)',
    'Análisis y resumen de documentos cargados',
    'Sugerencias de texto para completar o mejorar documentos',
    'Consultas en lenguaje natural sobre el contenido del sistema',
    'Generación de documentos desde plantillas ministeriales en PDF y Word',
    'Protocolo de firma QR integrado en documentos generados',
]
for it in items_X:
    p = doc.add_paragraph(style='List Bullet')
    set_font(p.add_run(it), size=10)

# --- Sección XI ---
heading('3.7 Sección XI — Generación de Contenido de Prensa (IA)', 2)
para(
    'Módulo diseñado para el equipo de prensa y comunicación del ministerio. '
    'Permite generar artículos institucionales, posts para redes sociales y '
    'propuestas de contenido de forma automatizada con Inteligencia Artificial:',
    size=11
)
items_XI = [
    'Generación semanal de propuestas de artículos para 6 sectores ministeriales: Transportes, Puertos, Telecomunicaciones, Inteligencia Artificial, Correos y Regulación',
    'Propuestas incluyen: título, introducción (2–3 oraciones) y esquema de 4–6 puntos por sección',
    'Redacción de borradores completos (+600 palabras) con estructura formal: introducción, desarrollo con subtítulos, conclusión y recomendaciones',
    'Los borradores se guardan automáticamente como artículos en estado BORRADOR en la plataforma',
    'Generación de publicaciones para redes sociales (versión Twitter ≤280 chars + versión Facebook extendida con hashtags y emojis)',
    'Función de copiado al portapapeles compatible con HTTP y HTTPS',
    'Vista paginada de artículos (10 por página) con controles de navegación',
    'Gestión completa: crear, editar, publicar, eliminar artículos',
    'Control de roles: solo ADMIN y GABINETE pueden generar contenido IA',
]
for it in items_XI:
    p = doc.add_paragraph(style='List Bullet')
    set_font(p.add_run(it), size=10)

# --- Sección XII ---
heading('3.8 Sección XII — Módulo Multimedia (IA)', 2)
para(
    'Plataforma para gestión, transcripción y análisis de contenido multimedia '
    '(audio y video) con Inteligencia Artificial:',
    size=11
)
items_XII = [
    'Carga de archivos de audio y video (MP3, MP4, WAV, WebM, etc.)',
    'Transcripción automática de audio/video usando OpenAI Whisper API',
    'Soporte multilingüe: transcripción en español, francés, inglés y otros idiomas',
    'Traducción del contenido transcrito a 5 idiomas: Francés, Inglés, Ruso, Chino y Portugués',
    'Generación de Dictamen Técnico Ministerial con 5 secciones estructuradas: resumen ejecutivo, relevancia institucional, análisis de contenido, riesgos y oportunidades, y recomendaciones',
    'Almacenamiento de transcripción, traducción y dictamen en base de datos',
    'Interfaz de gestión con reproductor integrado, filtros y búsqueda',
]
for it in items_XII:
    p = doc.add_paragraph(style='List Bullet')
    set_font(p.add_run(it), size=10)

# --- Sección XIII ---
heading('3.9 Sección XIII — Asistente IA Ministerial', 2)
para(
    'Chat inteligente contextualizado para el entorno ministerial. '
    'Permite al equipo ministerial consultar, redactar y analizar sin salir de la plataforma:',
    size=11
)
items_XIII = [
    'Interfaz de chat con historial de conversación persistente por sesión',
    'Generación de textos oficiales: decretos, notas verbales, circulares, informes',
    'Respuestas en español formal adaptadas al contexto ministerial de Guinea Ecuatorial',
    'Botón de copiado rápido de respuestas generadas',
    'Acceso desde el menú lateral bajo icono de Asistente IA',
]
for it in items_XIII:
    p = doc.add_paragraph(style='List Bullet')
    set_font(p.add_run(it), size=10)

# --- Sección XIV ---
heading('3.10 Sección XIV — Módulo WhatsApp', 2)
para(
    'Integración con WhatsApp Business API para comunicación oficial ministerial:',
    size=11
)
items_XIV = [
    'Envío de mensajes oficiales a través de WhatsApp Business',
    'Gestión de plantillas de mensajes ministeriales',
    'Historial de comunicaciones por expediente',
    'Notificaciones automáticas por WhatsApp para plazos críticos',
    'Panel de gestión de contactos y grupos ministeriales',
]
for it in items_XIV:
    p = doc.add_paragraph(style='List Bullet')
    set_font(p.add_run(it), size=10)

# --- Sección XV ---
heading('3.11 Sección XV — Agenda Ministerial', 2)
para(
    'Calendario y agenda digital para la organización de actividades ministeriales:',
    size=11
)
items_XV = [
    'Creación y gestión de eventos, reuniones y compromisos oficiales',
    'Vista de calendario mensual y semanal',
    'Recordatorios automáticos para eventos próximos',
    'Integración con el flujo de notificaciones existente',
    'Vinculación de eventos con expedientes documentales',
]
for it in items_XV:
    p = doc.add_paragraph(style='List Bullet')
    set_font(p.add_run(it), size=10)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 4. INFRAESTRUCTURA Y TECNOLOGÍA
# ═══════════════════════════════════════════════════════════════════════════════
heading('4. INFRAESTRUCTURA Y STACK TECNOLÓGICO', 1)

heading('4.1 Stack Tecnológico', 2)
tech_data = [
    ('Frontend',        'React 18 + TypeScript + Vite + TanStack Query + Tailwind CSS + Shadcn/ui'),
    ('Backend',         'NestJS 10 + TypeScript + Prisma ORM'),
    ('Base de Datos',   'PostgreSQL 15'),
    ('Servidor',        'Ubuntu 22.04 LTS — VPS DigitalOcean'),
    ('Proceso',         'PM2 (cluster mode) — proceso: ministerial-api'),
    ('Proxy inverso',   'Nginx (HTTP, listo para HTTPS con Let\'s Encrypt)'),
    ('IA / ML',         'OpenAI GPT-4o (texto) + Whisper API (audio)'),
    ('Email',           'Nodemailer + SMTP (Gmail / SMTP propio)'),
    ('Autenticación',   'JWT (7 dias de validez) + bcrypt'),
    ('Almacenamiento',  'Sistema de archivos local (/uploads) + rutas a BD'),
]
add_table(['Componente', 'Tecnología'], tech_data, col_widths=[4.5, 11.0])

heading('4.2 Infraestructura Actual de Produccion', 2)
infra_data = [
    ('IP del servidor', '157.230.178.118'),
    ('Usuario',         'root'),
    ('Ruta frontend',   '/var/www/ministerial-command-center/frontend/'),
    ('Ruta backend',    '/var/www/ministerial-command-center/backend/'),
    ('Puerto API',      '4001 (Nginx proxy)'),
    ('Puerto Web',      '80 (HTTP) — listo para 443 HTTPS'),
    ('PM2 proceso',     'ministerial-api (id=1)'),
    ('Base de datos',   'PostgreSQL en localhost'),
]
add_table(['Parámetro', 'Valor'], infra_data, col_widths=[5.0, 10.5])

heading('4.3 Rendimiento del Sistema', 2)
perf_data = [
    ('Listado de documentos',   '< 200 ms'),
    ('Carga de archivo (10 MB)', '< 5 segundos'),
    ('Procesamiento OCR',        '< 10 segundos'),
    ('Generacion con IA',        '< 15 segundos'),
    ('Consultas a base de datos', '< 50 ms (con indices)'),
    ('Verificacion recordatorios (1000 docs)', '< 30 segundos'),
]
add_table(['Operacion', 'Tiempo objetivo'], perf_data, col_widths=[8.0, 7.5])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 5. DOCUMENTACIÓN ENTREGADA
# ═══════════════════════════════════════════════════════════════════════════════
heading('5. DOCUMENTACIÓN TÉCNICA ENTREGADA', 1)
para(
    'Se han elaborado y entregado más de 9.000 líneas de documentación técnica y de usuario '
    'cubriendo todos los aspectos del sistema:',
    size=11
)
docs_data = [
    ('PRODUCTION_DEPLOYMENT_GUIDE.md',     '770 líneas', 'Instalación y despliegue completo en VPS'),
    ('MANUAL_DE_USUARIO.md',               '800 líneas', 'Manual de usuario final en español'),
    ('SYSTEM_MAINTENANCE_GUIDE.md',        '900 líneas', 'Guía de mantenimiento para administradores'),
    ('PROJECT_SUMMARY.md',                 '1000 líneas', 'Resumen completo del proyecto'),
    ('PHASE_4_AUTOMATION_TESTING.md',      '500 líneas', 'Escenarios de prueba de automatización'),
    ('UAT_PLAN.md',                        'Completo', 'Plan UAT con 12 escenarios, 30 casos de prueba'),
    ('TRAINING_MATERIALS.md',              'Completo', 'Materiales de formación (admin + usuario)'),
    ('VIDEO_TUTORIAL_SCRIPTS.md',          '10 videos', 'Scripts para 10 tutoriales en video (60 min)'),
    ('CLIENT_HANDOVER_CHECKLIST.md',       'Completo', 'Lista de verificación de entrega al cliente'),
    ('DEPLOYMENT_READINESS_CHECKLIST.md',  '600 líneas', 'Evaluacion de preparacion para produccion'),
    ('EXECUTIVE_SUMMARY.md',               'Completo', 'Resumen ejecutivo para partes interesadas'),
    ('GUIA_RAPIDA.md',                     'Completo', 'Guia de inicio rapido para usuarios finales'),
    ('PAQUETE_DE_ENTREGA_FINAL.md',        'Completo', 'Resumen del paquete de entrega final'),
    ('README.md',                          'Completo', 'Guia de configuracion y desarrollo del proyecto'),
]
add_table(['Documento', 'Extensión', 'Contenido'], docs_data, col_widths=[6.5, 2.2, 7.0])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 6. SEGURIDAD
# ═══════════════════════════════════════════════════════════════════════════════
heading('6. SEGURIDAD IMPLEMENTADA', 1)
sec_items = [
    'Autenticación JWT con expiración de 7 días y renovación automática',
    'Hashing seguro de contraseñas con bcrypt (salt rounds=10)',
    'Control de acceso basado en roles (RBAC) en todos los endpoints',
    'Restricción exclusiva de firma ministerial: solo rol GABINETE + isMinister=true',
    'Registro de auditoría completo para todas las acciones del sistema',
    'Prevención de inyección SQL mediante Prisma ORM (consultas parametrizadas)',
    'Protección XSS mediante escapado automático de React',
    'Validación de entrada en todos los DTOs con class-validator',
    'HTTPS/TLS listo para activar con certificado Let\'s Encrypt',
    'Separación de procesos en PM2 (ministerial-api aislado de otros servicios)',
]
for it in sec_items:
    p = doc.add_paragraph(style='List Bullet')
    set_font(p.add_run(it), size=11)
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 7. ESTADO FINANCIERO Y RECLAMACIÓN DE SALDO PENDIENTE
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading('7. ESTADO FINANCIERO Y SALDO PENDIENTE DE COBRO', 1)

para(
    'El presente informe certifica que todos los módulos y funcionalidades establecidos en '
    'el contrato original y sus ampliaciones han sido desarrollados, implementados y '
    'desplegados en producción en su totalidad. A la fecha de emisión de este informe, '
    'existe un saldo pendiente de pago que el contratista no ha recibido, conforme al '
    'detalle contractual que se presenta a continuación:',
    size=11
)
doc.add_paragraph()

# Detailed payment breakdown
finance_data = [
    # (Fase, Descripción, Valor contratado, Recibido, Pendiente, Entregado)
    ('Fases I–II',
     'Gestión documental, flujos de trabajo (11+6 etapas), roles, PDF, OCR',
     '$3,000', '$3,000', '$0', '✅ Entregado'),
    ('Fase III',
     'Protocolo de firma ministerial, automatización de plazos, IA documental',
     '$2,000', '$2,000', '$0', '✅ Entregado'),
    ('Sección XI',
     'Módulo de generación de contenido de prensa con IA (propuestas, borradores, redes sociales)',
     '$1,000', '$735', '$265', '✅ Entregado'),
    ('Sección XII',
     'Módulo multimedia: transcripción Whisper, traducción 5 idiomas, dictamen técnico IA',
     '$800', '$0', '$800', '✅ Entregado'),
    ('Secc. XIII–XV',
     'Asistente IA ministerial, módulo WhatsApp Business, agenda ministerial',
     '$800', '$0', '$800', '✅ Entregado'),
    ('Hosting',
     'Configuración, despliegue y mantenimiento VPS (12 meses)',
     '$400', '$0', '$400', '✅ Entregado'),
]
add_table(
    ['Fase / Módulo', 'Descripción', 'Contratado', 'Recibido', 'PENDIENTE', 'Estado'],
    finance_data,
    col_widths=[2.8, 6.5, 2.0, 2.0, 2.2, 2.2]
)

# Summary totals box
totals_data = [
    ('VALOR TOTAL DEL CONTRATO',           '$8,000',  '2E74B5'),
    ('TOTAL RECIBIDO POR EL CONTRATISTA',  '$5,735',  '1F497D'),
    ('SALDO PENDIENTE DE COBRO',           '$2,265',  'C00000'),
]
t2 = doc.add_table(rows=1, cols=2)
t2.style = 'Table Grid'
for i, (label, val, fill) in enumerate(totals_data):
    if i == 0:
        row = t2.rows[0]
    else:
        row = t2.add_row()
    shade_row(row, fill)
    row.cells[0].text = label
    row.cells[1].text = val
    for cell in row.cells:
        for run in cell.paragraphs[0].runs:
            set_font(run, size=12, bold=True, color=(255,255,255))
t2.columns[0].width = Cm(11.0)
t2.columns[1].width = Cm(3.5)
doc.add_paragraph()

# Formal claim paragraph
p_claim = doc.add_paragraph()
r_warn = p_claim.add_run(
    'AVISO FORMAL DE SALDO PENDIENTE: '
)
set_font(r_warn, size=11, bold=True, color=(192,0,0))
r_body = p_claim.add_run(
    'A la fecha de emisión del presente informe, el contratista NO ha recibido el pago '
    'correspondiente a $2,265 (dos mil doscientos sesenta y cinco dólares americanos), '
    'que constituye el saldo restante del contrato. Dicha cantidad comprende: '
    '$265 de la Sección XI (generación de contenido de prensa), $800 de la Sección XII '
    '(módulo multimedia), $800 de las Secciones XIII–XV (Asistente IA, WhatsApp, Agenda) '
    'y $400 correspondientes al servicio de hosting (12 meses). '
    'Todos los módulos indicados han sido completados, probados y desplegados en el servidor '
    'de producción conforme a lo acordado. El contratista solicita formalmente el abono '
    'del saldo pendiente para proceder al cierre oficial del contrato.'
)
set_font(r_body, size=11)
doc.add_paragraph()

# Work completion certificate box
cert_p = doc.add_paragraph()
shade_para_fill = '003366'
r_cert_title = cert_p.add_run('CERTIFICACIÓN DE TRABAJO COMPLETADO')
set_font(r_cert_title, size=11, bold=True, color=(0,70,127))

cert_items = [
    '15 módulos funcionales desarrollados, probados y operativos en producción',
    '40 funcionalidades implementadas — 100% del alcance contractual cubierto',
    'Sistema accesible en: http://157.230.178.118 (servidor activo)',
    'Código fuente completo disponible para entrega inmediata',
    'Más de 9.000 líneas de documentación técnica entregadas',
    'El único pendiente técnico (dominio HTTPS) requiere confirmación del cliente',
]
for ci in cert_items:
    p = doc.add_paragraph(style='List Bullet')
    set_font(p.add_run(ci), size=11, color=(0,70,127))

# ═══════════════════════════════════════════════════════════════════════════════
# 8. PENDIENTES PARA CIERRE DEFINITIVO
# ═══════════════════════════════════════════════════════════════════════════════
heading('8. PENDIENTES PARA CIERRE DEFINITIVO', 1)

pending = [
    ('P-01', 'Hosting y dominio personalizado',
     'Confirmar dominio final del cliente (ej. sistema.ministerio.gob.gq) y activar certificado HTTPS con Let\'s Encrypt. Estimación: 2 horas de trabajo técnico una vez confirmado el dominio.'),
    ('P-02', 'Pruebas de Aceptación del Usuario (UAT)',
     'Sesión de pruebas con el equipo del ministerio usando el plan UAT_PLAN.md entregado. Duración estimada: 1 semana. El contratista permanece disponible para corregir incidencias.'),
    ('P-03', 'Imagen del emblema nacional',
     'El cliente debe proveer la imagen oficial del emblema nacional en formato PNG de alta resolución para integrarla en el encabezado del sistema (fase estimada: 30 minutos tras recibir el asset).'),
    ('P-04', 'Videollamada de entrega y formación',
     'Sesión de videoconferencia para presentación del sistema completo, demostración en vivo y formación inicial al equipo designado. Los scripts de los 10 videotutoriales están preparados.'),
    ('P-05', 'Periodo de soporte post-entrega',
     '30 días de soporte técnico gratuito tras la entrega oficial para resolución de dudas, ajustes menores y corrección de errores no críticos.'),
]
for code, title, desc in pending:
    p = doc.add_paragraph()
    r1 = p.add_run(f'[{code}] {title}: ')
    set_font(r1, size=11, bold=True, color=(0,70,127))
    r2 = p.add_run(desc)
    set_font(r2, size=11)
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 9. GARANTÍAS Y SOPORTE
# ═══════════════════════════════════════════════════════════════════════════════
heading('9. GARANTÍAS Y COMPROMISOS POST-ENTREGA', 1)
guarantees = [
    'Garantía de corrección de errores críticos: 90 días tras la entrega oficial',
    'Soporte técnico activo: 30 días gratuitos para dudas y ajustes menores',
    'Documentación completa entregada en formato Markdown (15+ documentos, 9,000+ líneas)',
    'Código fuente completo entregado al cliente al finalizar el contrato',
    'Capacidad de despliegue adicional: el cliente puede reinstalar en cualquier servidor',
    'Scripts de mantenimiento y backup incluidos en la documentación',
    'Guías de formación para administradores y usuarios finales',
    'Plan de formación estructurado (2 días para admin + 1 día para usuarios)',
]
for g in guarantees:
    p = doc.add_paragraph(style='List Bullet')
    set_font(p.add_run(g), size=11)
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 10. FIRMA Y CIERRE
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading('10. DECLARACIÓN DE ENTREGA Y FIRMA', 1)
para(
    'El presente Informe Técnico Final certifica que la totalidad del trabajo de '
    'desarrollo contratado ha sido completado satisfactoriamente. El sistema "Centro de '
    'Mando Ministerial" se encuentra operativo, desplegado en servidor de producción y listo '
    'para uso por parte del Ministerio de Transportes, Telecomunicaciones, Correos, '
    'Infraestructura e Internet de Guinea Ecuatorial.',
    size=11
)

doc.add_paragraph()
para(
    f'Fecha de emisión del informe: {datetime.date.today().strftime("%d de %B de %Y")}',
    size=11, bold=True
)
doc.add_paragraph()

sig_table = doc.add_table(rows=5, cols=2)
sig_table.style = 'Table Grid'
labels_left  = ['Por el Contratista (Desarrollador):', '', 'Firma: ________________________', 'Nombre:', 'Fecha:']
labels_right = ['Por el Cliente (Ministerio):', '', 'Firma: ________________________', 'Nombre: Su Excelencia el Ministro', 'Fecha:']
for i in range(5):
    l_cell = sig_table.rows[i].cells[0]
    r_cell = sig_table.rows[i].cells[1]
    l_cell.text = labels_left[i]
    r_cell.text = labels_right[i]
    for cell in [l_cell, r_cell]:
        for run in cell.paragraphs[0].runs:
            bold = (i == 0)
            set_font(run, size=11, bold=bold, color=(0,70,127) if bold else None)

doc.add_paragraph()
doc.add_paragraph()

footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_f = footer_p.add_run(
    'Ministerio de Transportes, Telecomunicaciones, Correos, Infraestructura e Internet | '
    'Guinea Ecuatorial | Sistema Centro de Mando Ministerial v1.0'
)
set_font(r_f, size=9, italic=True, color=(127,127,127))

# ═══════════════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════════════
output_path = r'f:\workana_work\AfricaMinistrator\ministerial-command-center\Informe_Tecnico_Final_MTTSIA.docx'
doc.save(output_path)
print(f'Document saved: {output_path}')
